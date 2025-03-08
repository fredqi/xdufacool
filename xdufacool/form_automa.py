# form_automa.py ---
#
# Filename: form_automa.py
# Author: Fred Qi
# Created: 2022-08-21 16:02:49(+0800)
#
# Last-Updated: 2023-11-06 15:01:34(+0800) [by Fred Qi]
#     Update #: 2444
# 

# Commentary:
#
#
# 

# Change Log:
#
#
# 
import re
import os
# import sys
# import fitz
import logging
import warnings
# import markdown
# import frontmatter
from os import path
from glob import glob
from pathlib import Path
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
# from dataclasses import dataclass, field, asdict, fields
from configparser import ConfigParser
# from configparser import ExtendedInterpolation

from xdufacool.utils import format_list
from xdufacool.score_helper import ScoreStat
from xdufacool.score_helper import ScoreAnalysis

import openpyxl
from mailmerge import MailMerge
# from datetime import datetime
# from PyPDF2 import PdfReader
# from PyPDF2 import PdfMerger
# from docx import Document
import docx
from docx.styles import style
from docx.shared import Cm, Pt
# from docx.text.parfmt import ParagraphFormat
# from docx.styles.style import _TableStyle
# from docx.styles.style import _ParagraphStyle
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

# pip install mistletoe-ebp
import mistletoe
# from mistletoe.block_tokens import Heading, Paragraph, HTMLBlock
from mistletoe.block_token import Heading, Paragraph, Table
# from mistletoe.block_token import Table
# , HTMLSpan
from mistletoe.span_token import RawText, LineBreak, Strong, Emphasis


# if 'linux' == sys.platform:
#     # WORKDIR = "/home/fred/cloud/share/senior-design"
#     # WORKDIR = "/home/fred/cloud/OneDrive/students/forms-undergraduates"
#     WORKDIR = "/home/fred/cloud/OneDrive/students/senior-design/templates"
#     TEACHDIR = "/home/fred/cloud/OneDrive/students/senior-design"
#     # TEACHDIR = "/home/fred/cloud/share/teaching"
# elif 'win32' == sys.platform:
#     from docx2pdf import convert
#     WORKDIR = "C:\\Users\\fredq\\github\\senior-design"
#     TEACHDIR = "C:\\Users\\fredq\\github\\teaching"

# DATA_FILE = path.join(TEACHDIR, "毕业设计表单数据-2024.xlsx")
# TEMPLATE_DIR = path.join(WORKDIR, "")
# OUTPUT_DIR = WORKDIR


class MarkdownDocxMerger:
    def __init__(self, config_path):
        self.config = self.load_config(config_path)
        self.templates_folder = self.config.get('Settings', 'templates_folder')
        self.templates = self.load_templates()
        self.sections = {}

    @staticmethod
    def get_text(token):
        if isinstance(token, RawText):
            return token.content
        elif not isinstance(token, LineBreak):
            raw_text = token.children[0]
            assert isinstance(raw_text, RawText)
            return raw_text.content.strip()

    def load_config(self, config_path):
        config = ConfigParser()
        config.read(config_path)
        return config

    def load_templates(self):
        # Load templates from a specified folder
        templates = {}
        for template_name in self.config['Templates']:
            template_path = path.join(self.templates_folder,
                                      self.config.get('Templates', template_name))
            templates[template_name] = template_path
        return templates

    def parse_markdown(self, markdown_path):
        # with open(markdown_path, 'r', encoding='utf-8') as md_file:
        with open(markdown_path, 'r') as istream:
            text = istream.read()
            doc = mistletoe.Document.read(text, front_matter=True)
            metadata = doc.front_matter.get_data()
            sections = {}
            sec_title, key = 'section', 'default'
            for item in doc.children:
                if isinstance(item, Heading):
                    if 1 == item.level:
                        sec_title = self.get_text(item)
                        sections[sec_title] = {}
                    elif 2 == item.level:
                        key = self.get_text(item)
                        # print(key)
                        sections[sec_title][key] = []                        
                else:
                    sections[sec_title][key].append(self.get_text(item))                    
        # Return a dictionary containing both metadata and sections
        return metadata, sections

    def generate_docx(self, markdown_path):
        metadata, sections = self.parse_markdown(markdown_path)
            
        for sec_title, content in sections.items():
            # print(sec_title, content)
            if sec_title in self.templates:
                data = {key: value for key, value in metadata.items()}
                data.update(content)
                self.merge_and_create_docx(sec_title, data)

    def merge_and_create_docx(self, sec_title, data):
        # Perform the mail merge
        template_path = self.templates[sec_title]
        merge_data = {}
        for key, value in data.items():
            field = self.config.get('Mappings', key)
            if 'date' == field:
                merge_data['year'] = str(value.year)
                merge_data['year_p'] = str(value.year - 1)
            if isinstance(value, list):
                merge_data[field] = '\n'.join(value)
            else:
                merge_data[field] = str(value)
            
        with MailMerge(template_path) as document:
            missing_fields = document.get_merge_fields() - merge_data.keys()
            if missing_fields:
                print(sec_title, missing_fields)
            document.merge(**merge_data)
            # Optionally, handle the content if needed to be inserted into the document
            output_filename = f"{sec_title}-{data.get('姓名', '姓名')}.docx"
            document.write(output_filename)
            print(f"Generated file: {output_filename}")


class Markdown2Docx(object):
    # def __init__(self, score_stat, score_analysis, filename, styles):
    def __init__(self, summary_markdown, styles):
        # help(Heading)
        tag_start = r'<(?P<name>\w+)\s+class="(?P<html_class>\w+)">'
        tag_end = r'</\w+>'
        table_pattern = tag_start + r'\s*' + tag_end
        self.tag_parser = re.compile(table_pattern)
        self.tag_start = re.compile(tag_start)
        self.tag_end = re.compile(tag_end)
        # self.score_stat = score_stat
        # self.score_analysis = score_analysis
        self.styles = styles
        self.sections = {}
        # with open(summary_markdown, 'r') as istream:
        logging.debug("    Initializing Markdown2Docx ...")
        doc = mistletoe.Document(summary_markdown)
        for item in doc.children:
            if isinstance(item, Heading) and 1 == item.level:
                key = self.get_text(item)
                contents = list()
                self.sections[key] = contents
            else:
                contents.append(item)

    @staticmethod
    def get_text(token):
        if isinstance(token, RawText):
            return token.content
        elif not isinstance(token, LineBreak):
            raw_text = token.children[0]
            assert isinstance(raw_text, RawText)
            return raw_text.content.strip()

    def add_exam_stat_table(self, cell):
        """Add a table containing statistics of exam scores."""
        headers = ['课序号', '人数'] + ['人数', '百分比']*5 + ['平均分']
        table = cell.add_table(rows=3, cols=len(headers))
        # table.style = self.styles['Simple Table']
        table.style = self.styles['Score Table']
        row_cells = table.rows[1].cells
        for idx, text in enumerate(headers):
            row_cells[idx].text = text

        intervals = self.score_stat.stat
        headers = [f'{iv.left}-{iv.right}' for iv in intervals[:-1]]
        headers.append(intervals[-1].desc)
        row_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            index = 2*idx + 2
            row_cells[index].text = text

        status = [self.score_stat.course_order,
                  f'{self.score_stat.n_students:d}']
        for iv in intervals:
            status.extend([f'{iv.count:d}', f'{iv.percent:4.1f}%'])
        status.append(f'{self.score_stat.average:5.2f}')
        row_cells = table.rows[2].cells
        for idx, text in enumerate(status):
            row_cells[idx].text = text

        # merge cells
        for col in [0, 1, 12]:
            cell_a, cell_b = table.cell(0, col), table.cell(1, col)
            cell_a.merge(cell_b)
        for col in range(2, 12, 2):
            cell_a, cell_b = table.cell(0, col), table.cell(0, col + 1)
            cell_a.merge(cell_b)

    def add_content(self, cell, contents):
        for idx, item in enumerate(contents):
            if 0 == idx:
                tblp = cell.paragraphs[0]
                if "Heading 2" in self.styles:
                    tblp.style = self.styles["Heading 2"]
            else:
                tblp = cell.add_paragraph()
            if isinstance(item, Heading):
                tblp.style = self.styles["Heading 2"]
                run = tblp.add_run()
                run.add_text(self.get_text(item)).bold = True
            elif isinstance(item, Paragraph):
                for item in item.children:
                    if isinstance(item, LineBreak):
                        tblp = cell.add_paragraph()
                        continue
                    tblp.style = self.styles["Table Text"]
                    run = tblp.add_run()
                    content = self.get_text(item)
                    if isinstance(item, Emphasis):
                        content = f"  {content}  "
                    run.add_text(content)
                    if isinstance(item, Strong):
                        run.bold = True
                    elif isinstance(item, Emphasis):
                        run.underline = True
            elif isinstance(item, Table):
                self.add_table(cell, item)
            # elif isinstance(item, HTMLBlock):
            #     m = self.tag_parser.match(item.children)
            #     if m:
            #         tag_name, tag_class = m.groups()
            #         if "table" == tag_name and "score_stat" == tag_class:
            #             # print("To add a table containing score statistics.")
            #             # table = cell.add_table(2, 2)
            #             self.add_exam_stat_table(cell)
            else:
                print("Unhandled", type(item), item.content)

    def add_table_row(self, row_cells, items, header=False):
        """Add a row to a table."""
        for idx, cell_md in enumerate(items):
            contents = cell_md.children
            p = row_cells[idx].paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item in contents:
                if isinstance(item, RawText):
                    text = self.get_text(item)
                    run = p.add_run(text)
                    run.font.size = Pt(12)
                    if header:
                        run.bold = True
                elif isinstance(item, HTMLSpan):
                    m_start = self.tag_start.match(item.content)
                    m_end = self.tag_start.match(item.content)
                    logging.debug("HTMLSpan", item.content, m_start, m_end)
                    if m_start:
                        tag_name, tag_class = m_start.groups()
                        if "span" == tag_name:
                            text = self.score_analysis.get_text(tag_class)
                            # print(tag_class, text)
                            run = p.add_run(text)
                            run.font.size = Pt(12)
                    elif m_end:
                        tag_name = m_end.group('name')
                        if 'span' == tag_name:
                            continue
                else:
                    print(type(item), item)
        
    def add_table(self, cell, table_md):
        """Add a table inside a table cell."""
        # n_rows = len(table_md.children)+1
        n_cols = len(table_md.header.children)
        table = cell.add_table(rows=0, cols=n_cols)
        # table.style = self.styles['Simple Table']
        table.style = self.styles['Score Table']
        table.autofit = True
        table.allow_autofit = True
        self.add_table_row(table.add_row().cells,
                           table_md.header.children,
                           header=True)
        for row_md in table_md.children:
            self.add_table_row(table.add_row().cells,
                               row_md.children)


class RowMapper(object):
    def __init__(self, keys, headers):
        self.indices = []
        # logging.debug(f"RowMapper: {keys} {headers}")
        for key in keys:
            self.indices.append(headers.index(key))

    def get(self, values):
        return [values[idx] for idx in self.indices]                              
            

def mailmerge_fields(sheet):
    """Get mail merge fields from a given excel sheet."""
    fields = []
    for row in sheet.iter_rows(min_row=2, max_col=2):
        fields.append((c.value for c in row))
    return dict(fields)


def load_form_data(filename):
    """Load data from excel file."""
    wb = openpyxl.load_workbook(filename, data_only=True)
    fields = mailmerge_fields(wb["Fields"])
    return fields, wb


def sheet_column_keys(sheet, fields):
    """Create dict keys from an excel sheet."""
    # print(sheet.title)    
    header = sheet.iter_rows(max_row=1, values_only=True)
    keys = [fields[value] for value in list(header)[0]]
    return keys


def sheet_row_dict(sheet, keys, min_row=4, max_row=8):
    """Create a dictionary from a row of a excel sheet."""
    score_range = {'及格': 60, '中等': 70, '良好': 80, '优秀': 90}
    data = []
    for row in sheet.iter_rows(min_row=min_row, max_row=max_row):
        row_dict = {k:v.value for k, v in zip(keys, row)}
        if 'sc' in row_dict and 'score' in row_dict:
            # print(row_dict['sc'], row_dict['score'])
            sc, score = row_dict['sc'], row_dict['score']
            sc = 0 if sc is None else sc
            sc_min = score_range[score]
            # print(sc, score, sc_min)
            if sc < sc_min or sc >= (sc_min + 10):
                warnings.warn(f"成绩等级与分数不一至 {row_dict['id']} {sc} {score}")
                
        if 'date' in row_dict:
            dt = row_dict['date']
            row_dict['date_data'] = dt
            row_dict['date'] = f'{dt.year}年{dt.month}月{dt.day}日'
            row_dict['year'] = f'{dt.year}'
            row_dict['year_p'] = f'{dt.year-1}'
        data.append(row_dict)
    return data


def document_merge(template, data):
    """Merge data into a docx template."""
    output_filename = path.basename(template)
    print(output_filename)
    for row in data:
        document = MailMerge(template)
        missing_fields = document.get_merge_fields() - row.keys()
        if missing_fields:
            print(output_filename, missing_fields)
        document.merge(**row)
        sid, name = row['id'], row['name']
        title = row['title']
        output_dir = path.join(OUTPUT_DIR, f'{title}')
        if not path.exists(output_dir):
            os.mkdir(output_dir)
        # print(output_filename, sid, name, row['date'])
        # print(f"    {sid} {name} {row['sc']} {row['score']}")
        document.write(path.join(output_dir, output_filename))
           
    
def template_dict(template_dir, sheets):
    """Create a dictionary of templates."""
    sheet_names = list(sheets)
    template_files = glob(f"{template_dir}/*.docx")
    print(template_dir, "\n".join(template_files))
    templates = {}
    for filename in template_files:
        for name in sheet_names:
            if filename.find(name) >= 0:
                templates[name] = filename
                sheet_names.remove(name)
    return templates

    
def form_auto_merge(sheets=['日常考核']):
    # sheet_names = ['日常考核', '软硬件验收', '中期检查', '指导教师评分表', '评阅评分表', '答辩登记表']
    templates = template_dict(TEMPLATE_DIR, sheets)
    print(sheets)
    print("\n".join([f'{k}={v}' for k,v in templates.items()]))
    fields, wb = load_form_data(DATA_FILE)
    # print(fields)
    for sheet_name in sheets:
        template = templates[sheet_name]
        keys = sheet_column_keys(wb[sheet_name], fields)
        data = sheet_row_dict(wb[sheet_name], keys)
        document_merge(template, data)


def merge_replacements(folder):
    sheet_names = ['日常考核', '中期检查', '指导教师评分表', '评阅评分表', '软硬件验收', '答辩登记表']
    files_pdf = glob(path.join(folder, "*.pdf"))
    merger = PdfMerger()

    for name in sheet_names:
        pdf_file = glob(path.join(folder, f"*{name}*.pdf"))
        pdf_stream = open(pdf_file[0], 'rb')
        merger.append(pdf_stream)

    output = open(path.join(folder,  "replacements.pdf"), 'wb')
    merger.write(output)
    merger.close()
    output.close()


def replace_pages(student_id, name, page_begin, page_end):
    filename = f"{student_id}-{name}-毕业设计归档材料.pdf"
    archive_filepath = path.join(WORKDIR, "archived", filename)
    
    reader = PdfReader(archive_filepath)
    n_pages = len(reader.pages)
    archive_stream = open(archive_filepath, 'rb')

    replacement_filepath = path.join(WORKDIR, f"{student_id}-{name}", "replacements.pdf")
    replacement_stream = open(replacement_filepath, 'rb')

    merger = PdfMerger()
    merger.append(archive_stream, pages=(0, page_begin))
    merger.append(replacement_stream)
    merger.append(archive_stream, pages=(page_end, n_pages))

    output_filepath = path.join(WORKDIR, "updated", filename)    
    output_stream = open(output_filepath, 'wb')
    merger.write(output_stream)

    merger.close()
    output_stream.close()    


def convert_merge_replace():
    folders = glob(path.join(WORKDIR, "1*"))
    for folder in folders:
        convert(folder)
        merge_replacements(folder)


def sheet_replace():
    wb = openpyxl.load_workbook(DATA_FILE, data_only=True)
    sheet = wb["replace"]
    for row in sheet.iter_rows(min_row=2, values_only=True):
        print(row)
        replace_pages(*row)


class SummaryComposer(object):
    """Composer for teaching summary."""

    def __init__(self, course_order, base_dir, workspace_dir):
        self.table_width = Cm(15)
        self.table_height = Cm(22.5)
        self.course_order = course_order
        self.document = None
        self.base_dir = base_dir
        self.summary_dir = workspace_dir / "summary"
        self.summary_dir.mkdir(parents=True, exist_ok=True)
        logging.debug(f"    Summary dir: {self.summary_dir.relative_to(self.base_dir)}")
        
    # def __init__(self, course, term):
    #     self.course = course
    #     self.term = term
    #     self.table_width = Cm(15)
    #     self.table_height = Cm(22.5)
    #     self.document = docx.Document(course.summary_filepath)
    #     self.setup_core_properties()
    #     self.config_styles()
    #     record_filepath = path.join(term.data_dir, term.record)
    #     wb = openpyxl.load_workbook(record_filepath)
    #     self.teaching_records = wb.active
    #     self.load_scores(course.score_filepath)
    #     text_filepath = path.join(term.data_dir, term.summary_text)
    #     self.md2docx = Markdown2Docx(self.score_stat,
    #                                  self.score_analysis,
    #                                  text_filepath, self.styles)
        
    @staticmethod
    def set_center(cell):
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def set_bold(cell, text):
        p = cell.paragraphs[0]
        p.add_run(text).bold = True

    # @staticmethod
    # def display(fmt):
    #     if isinstance(fmt, _ParagraphStyle):
    #         print(fmt.name)
    #         print("  ", fmt.paragraph_format.alignment)
    #         print("  ", fmt.font.name, fmt.font.size, fmt.font.bold)
    #         print("  ", fmt.paragraph_format.line_spacing)
    #         print("  ", fmt.paragraph_format.left_indent)
    #         print("  ", fmt.paragraph_format.first_line_indent)

    def config_styles(self):
        self.styles = {}
        for item in self.document.styles:
            name = item.name
            if name.find("Table") >= 0:
                self.styles[name] = item
            elif name.find("Heading") >= 0:
                self.styles[name] = item
        # # print("styles:", self.styles)
        # for key, style in self.styles.items():
        #     if key.find("Table") >= 0:
        #         print(key, style)
                    
    def setup_core_properties(self):
        """Setup core properties of the summary file."""
        core_prop = self.document.core_properties
        core_prop.author = self.course.teachers
        core_prop.created = self.course.date
        core_prop.last_printed = self.course.date
        core_prop.title = f"{self.course.course_name} 教学一览表"
        core_prop.subject = f"{self.course.course_name} {self.course.course_id}"
        core_prop.comments = str(self.course)

    def load_scores(self, score_filepath):
        sheet = openpyxl.load_workbook(score_filepath).active
        # TODO: Validate the sheet title by semester and course order
        # sheet_title = f"{self.course.semester}-{self.course.course_order}"
        #assert sheet_title == sheet.title

        self.scores = []
        for row in sheet.iter_rows(min_row=5):
            self.scores.append([item.value for item in row])
        index = self.scores[0].index("期末成绩")
        # print(index, self.scores[0][index])
        scores_exam = [row[index] for row in self.scores[1:]]
        self.score_stat = ScoreStat(scores_exam)
        if '实验成绩' in self.scores[0]:
            index = self.scores[0].index("实验成绩")
        else:
            index = self.scores[0].index("平时成绩")
        scores_hw = [row[index] for row in self.scores[1:]]
        self.score_analysis = ScoreAnalysis()
        self.score_analysis.add_parts(['homeworks', 'final'],
                                      [scores_hw, scores_exam])

    def add_table(self, heading, rows, cols=1,
                  headers=None, widths=None, style=None):
        self.document.add_page_break()
        self.document.add_heading(heading, 1)
        n_cols = len(headers) if headers else cols
        table_style = self.styles.get(style, None)
        table = self.document.add_table(rows=rows, cols=n_cols,
                                        style=table_style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if headers and widths and len(headers) == len(widths):
            table.autofit = False
            table.allow_autofit = False
            for idx in range(len(widths)):
                table.columns[idx].width = widths[idx]
                table.cell(0, idx).width = widths[idx]
        else:
            table.autofit = True
            table.allow_autofit = True

        # framed page contents
        if 1 == rows and 1 == n_cols:
            table.columns[0].width = self.table_width
            table.rows[0].height = self.table_height
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].style = self.styles["Heading 2"]
        elif headers:
            for idx, text in enumerate(headers):
                cell = table.cell(0, idx)
                p = cell.paragraphs[0]
                p.add_run(text).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return table

    def add_teaching_record(self):
        """Add a table recording the teaching process."""
        records = self.teaching_records
        keys = [f"上课时间-{self.course_order}", "知识点", "教学过程记录"]
        row_mapper = RowMapper(keys, next(records.values))
        
        headers = ["序号", "上课时间", "知识点", "教学过程记录"]
        widths = [Cm(1.25), Cm(2.5), Cm(4.25), Cm(7.5)]
        table = self.add_table("教学过程记录", rows=1,
                               headers=headers, widths=widths,                               
                               style="Teaching Table")

        idx = 1
        for row in records.iter_rows(min_row=2, values_only=True):
            values = row_mapper.get(row)
            row_cells = table.add_row().cells
            self.set_center(row_cells[0])
            self.set_bold(row_cells[0], f"{idx:d}")
            # rint(values, type(values[0]))            
            # logging.debug(f"{idx:d} {values[0]}")
            row_cells[1].text = f"{values[0]:%Y-%m-%d}"
            self.set_center(row_cells[1])
            
            row_cells[2].text = values[1]
            
            row_cells[3].paragraphs[0].text = values[2]
            p = row_cells[3].add_paragraph()
            p.text = "课堂秩序良好"
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1

    def add_framed_section(self, section_title):
        """For normal summary sections provided in markdown, which
        include the following contents:
          - 课程教学目标及与毕业要求的对应关系
          - 课程考核方式及成绩评定原则
          - 试卷（卷面）分析及总结
          - 课程能力达成度分析及总结
        """
        
        table = self.add_table(section_title, rows=1, cols=1,
                               style="Frame Table")
        
        cell = table.cell(0, 0)
        contents = self.md2docx.sections[section_title]
        self.md2docx.add_content(cell, contents)

    def add_score_table(self):
        """Add score table."""
        widths = [1.25, 2.5, 2, 1.25, 1.75, 1.9, 1.9, 1.9, 1.25]
        table = self.add_table("西安电子科技大学学习过程考核记录表",
                               rows=1, headers=self.scores[0],
                               widths=[Cm(w) for w in widths],
                               style="Score Table")
        for row in self.scores[1:]:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                text = "" if value is None else f"{value}"
                p = row_cells[idx].paragraphs[0]
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p.text = text

    def add_exam_table(self):
        """Add exam sheet and grading rules."""
        table = self.add_table("课程考试试题及答案", rows=1, cols=1,
                               style="Frame Table")
        
        run = table.cell(0, 0).paragraphs[0].add_run()
        run.add_text("（试题、标准答案及评分细则）").bold = True
        run = table.cell(0, 0).add_paragraph().add_run()

        # for image in sorted(self.term.images):
        #     run.add_picture(image)

    def fill_titlepage(self, student_group, summary_config):
        """Fills the title page of the summary document using a template."""
        # working_dir = Path(summary_config.get('working_dir', '.')) 
        # template = Path(summary_config.get('template', ""))
        template = self.base_dir / 'templates' / summary_config.get('template', '')
        logging.debug(f"    Using template: {template.relative_to(self.base_dir)}")
        if not template.exists():
            logging.error(f"! Summary template {template.relative_to(self.base_dir)} does not exist.")
            return
        teachers_all = student_group.course.teachers
        teachers = [teachers_all[key].name for key in student_group.teacher_ids]
        summary_date = datetime.strptime(summary_config['date'], "%Y-%m-%d")
        summary_date_desc = summary_date.strftime("%Y年%m月%d日")
        merger = MailMerge(template)
        title_info = {"semester": student_group.course.semester,
                      "course_name": student_group.course.topic,
                      "hours": student_group.course.teaching_hours,
                      "credits": student_group.course.credits,
                      "classes": ", ".join(student_group.admin_classes),
                      "teachers": format_list(teachers, conj="、", lang=student_group.course.language),
                      "summary_date": summary_date_desc}
        # logging.debug(f"    Title info: {title_info}")
        merger.merge(**title_info)

        parts = ["教学一览表", student_group.course.topic, student_group.course.semester, student_group.group_id]
        filename = '-'.join(parts) + ".docx"
        logging.info(f"* Creating summary {filename}")
        student_group.summary_filepath = self.summary_dir / filename
        merger.write(student_group.summary_filepath)

    def fill_score_analysis(self, score_filepath, summary_text):
        self.load_scores(score_filepath)
        env = Environment(loader=FileSystemLoader(summary_text.parent),
                          trim_blocks=True, lstrip_blocks=True)
        template = env.get_template(summary_text.name)
        score_table = self.score_stat.get_table()
        analysis = self.score_analysis.get_analysis()
        rendered_markdown = template.render(score_table=score_table, **analysis)
        return rendered_markdown

    def create_summary(self, summary_filepath, teaching_records, score_filepath, summary_text):
        self.document = docx.Document(summary_filepath)
        self.config_styles()
        wb = openpyxl.load_workbook(teaching_records)
        self.teaching_records = wb.active
        logging.debug(f"    Using summary text: {summary_text.relative_to(self.base_dir)}")
        logging.debug(f"    Using score file: {score_filepath.relative_to(self.base_dir)}")
        rendered_markdown = self.fill_score_analysis(score_filepath, summary_text)
        # print(rendered_markdown)
        self.md2docx = Markdown2Docx(rendered_markdown, self.styles)
        self.add_teaching_record()
        self.add_framed_section("课程教学目标及与毕业要求的对应关系")
        self.add_framed_section("课程考核方式及成绩评定原则")
        self.add_score_table()
        self.add_exam_table()
        self.add_framed_section("试卷（卷面）分析及总结")
        self.add_framed_section("课程能力达成度分析及总结")
        self.document.save(summary_filepath)

class PdfPage2Image(object):
    
    def __init__(self, dpi=300,
                 hmargin=3.5, vmargin=3.5,
                 paper_width=21, paper_height=29.7):
        self.dpi = dpi
        # 72 is the default sampling DPI
        # self.mat = fitz.Matrix(dpi/72, dpi/72)

        self.width  = paper_width - 2*hmargin
        self.height = paper_height - 2*vmargin
        left   = self.cm_to_dots(hmargin, self.dpi)
        top    = self.cm_to_dots(vmargin, self.dpi)
        right  = self.cm_to_dots(paper_width - hmargin, self.dpi)
        bottom = self.cm_to_dots(paper_height - vmargin, self.dpi)
        self.clip = fitz.Rect(left, top, right, bottom)

    @staticmethod
    def cm_to_dots(x, dpi=72):
        return x/2.54*dpi

    def convert(self, filepath, output_dir):
        if not path.exists(filepath):
            return None

        if not path.exists(output_dir):
            os.mkdir(output_dir)

        images = []
        doc = fitz.open(filepath)
        for page in doc:
            pixels = page.get_pixmap(dpi=self.dpi)
            image_clip = fitz.Pixmap(pixels, pixels.width, pixels.height, self.clip)
            image_clip.set_dpi(self.dpi, self.dpi)
            image_filepath = path.join(output_dir, f"page-{page.number:02d}.png")
            image_clip.save(image_filepath)
            images.append(image_filepath)

        return images

# @dataclass
# class TermInfo(object):
#     """Information of a term."""
    
#     semester: str
#     data_dir: str
#     record: str
#     exam: str
#     summary_text: str
#     working_dir: str
#     images: list = field(init=False)

#     def __str__(self):
#         return f"{self.semester}"

#     def __post_init__(self):
#         """Convert the exam pdf file to images."""
#         output_dir = path.join(self.working_dir, f'.{self.semester}')
#         exam_filepath = path.join(self.data_dir, self.exam)
#         # if path.exists(exam_filepath):
#         #     pdf2image = PdfPage2Image()
#         #     self.images = pdf2image.convert(exam_filepath, output_dir)


# @dataclass
# class CourseInfo(object):
#     """Information of a course."""

#     semester: str
#     course_id: str
#     course_name: str
#     course_order: int
#     credits: float
#     hours: int
#     teachers: str
#     classes: str
#     summary_date: str
#     data_dir: str
#     working_dir: str
#     date: datetime = field(init=False)
#     score_filepath: str = field(init=False)
#     summary_file: str = field(init=False)
#     summary_filepath: str = field(init=False)

#     def __post_init__(self):
#         self.date = datetime.strptime(self.summary_date, '%Y-%m-%d')
#         dt = self.date
#         self.summary_date = self.date.strftime(f'{dt.year}年{dt.month}月{dt.day}日')

#         score_file = f"scores-{self.course_order}.xlsx"
#         self.score_filepath = path.join(self.data_dir, score_file)
#         # Filename of the teaching summary.
#         parts = ["教学一览表", self.course_name, self.course_id,
#                  self.semester, self.course_order]
#         self.summary_file = '-'.join(parts) + ".docx"
#         self.summary_filepath = path.join(self.working_dir, self.summary_file)
        
#     def __str__(self):
#         return f"{self.course_id:8} {self.course_name} {self.semester} {self.course_order} {self.classes} {self.teachers}"

#     def compose_summary(self, template, term):
#         # Create the title page with mail merge.
#         merger = MailMerge(template)
#         merger.merge(**asdict(self))
#         merger.write(self.summary_filepath)

#         # Create the full summary file
#         composer = SummaryComposer(self, term)
#         composer.create_summary()
       

# def load_config(config_filepath):
#     config = ConfigParser(interpolation=ExtendedInterpolation())
#     config.read(config_filepath)
#     template = config.get('general', 'template_filepath')
    
#     fields_init = filter(lambda f: f.init, fields(CourseInfo))
#     course_keys = [f.name for f in fields_init]
#     fields_init = filter(lambda f: f.init, fields(TermInfo))
#     term_keys = [f.name for f in fields_init]
    
#     courses, terms = [], {}
#     data = dict(config.items('general', raw=False))
#     for cls in config['general']['classes'].split(','):
#         data.update(config.items(cls, raw=False))
#         term_key = data['term_key']
#         data.update(config.items(term_key, raw=False))

#         properties = {key: data[key] for key in course_keys}
#         info = CourseInfo(**properties)            
#         courses.append(info)
        
#         semester = data['semester']
#         if semester not in terms:
#             properties = {key: data[key] for key in term_keys}
#             terms[semester] = TermInfo(**properties)

#     return courses, terms, template
        
  
if __name__ == "__main__":

    # for config_filepath in glob("/home/fred/lectures/PRML/eval/summary/teaching-21+22.ini"):
    #     courses, terms, template_filepath = load_config(config_filepath)
    #     for course in courses:
    #         print(course)
    #         course.compose_summary(template_filepath, terms[course.semester])

    # sheets = ['日常考核', '软硬件验收', '中期检查', '指导教师评分表', '评阅评分表', '答辩登记表']
    sheets = ['任务书']
    form_auto_merge(sheets)
    # # convert_merge_replace()
    # # merge_replacements("16020520025-马丁扬")
    # # replace_pages("16020520025", "马丁扬", 17, 26)
 
    # attentions = []
    # for folder in folders:
    #     print(f"Processing {folder}...")
    #     convert(folder)
    #     pdfs = glob(f"{folder}/*.pdf")
    #     for pdf in pdfs:
    #         reader = PdfReader(pdf)
    #         n_pages = len(reader.pages)
    #         if n_pages > 1:
    #             print(n_pages, pdf)
    #             attentions.append(f"{n_pages} pages in {pdf}")
    # print("\n".join(attentions))
    
# 
# form_automa.py ends here
